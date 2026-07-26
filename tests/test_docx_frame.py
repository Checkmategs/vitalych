"""Smoke tests for DOCX frame presets."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.docx_frame import apply_frame
from src.md_to_docx import markdown_to_docx
from src.style_profile import load_style_profile


class DocxFrameTest(unittest.TestCase):
    def test_none_is_noop(self) -> None:
        doc = Document()
        doc.add_paragraph("Hello")
        warnings = apply_frame(doc, {"frame": {"preset": "none"}}, {})
        self.assertEqual(warnings, [])

    def test_frame_only_writes_docx(self) -> None:
        profile = load_style_profile(Path("style-profile.yaml"))
        profile = {**profile, "frame": {**(profile.get("frame") or {}), "preset": "frame_only"}}
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "t.docx"
            markdown_to_docx("# Title\n\nBody text.", out, profile=profile, data={"system": {"name": "АС"}})
            self.assertTrue(out.is_file())
            Document(str(out))  # opens

    def test_stamp_compact_writes_docx(self) -> None:
        profile = load_style_profile(Path("style-profile.yaml"))
        profile = {
            **profile,
            "frame": {
                "preset": "stamp_compact",
                "stamp_fields": {
                    "designation": "{{ system.topic_code }}",
                    "title": "{{ system.name }}",
                    "developer": "{{ parties.developer }}",
                },
            },
        }
        data = {
            "system": {"name": "Тест", "topic_code": "ABC.01"},
            "parties": {"developer": "ООО Разработчик"},
        }
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "t.docx"
            markdown_to_docx("# Title\n\nBody.", out, profile=profile, data=data)
            self.assertTrue(out.is_file())


if __name__ == "__main__":
    unittest.main()
