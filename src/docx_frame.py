"""Apply configurable GOST-like page frame / stamp to a python-docx Document."""

from __future__ import annotations

from typing import Any

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from jinja2 import BaseLoader, Environment, StrictUndefined, UndefinedError

from src.cases import make_case_filter, prepare_render_data

KNOWN_PRESETS = frozenset({"none", "frame_only", "stamp_compact", "eskd_2_2a"})


def _set_page_borders(section, size: str = "12", space: str = "24") -> None:
    sect_pr = section._sectPr
    # Remove existing pgBorders if present
    for child in list(sect_pr):
        if child.tag == qn("w:pgBorders"):
            sect_pr.remove(child)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), "000000")
        pg_borders.append(el)
    sect_pr.append(pg_borders)


def _render_stamp_field(template: str, data: dict[str, Any]) -> str:
    if not template or not str(template).strip():
        return ""
    wrapped, warnings = prepare_render_data(data)
    env = Environment(loader=BaseLoader(), undefined=StrictUndefined)
    env.filters["case"] = make_case_filter(warnings)
    try:
        return env.from_string(str(template)).render(**wrapped).strip()
    except (UndefinedError, Exception):
        # Soft failure for stamp cells — leave empty rather than failing whole DOCX
        return ""


def _add_compact_stamp(doc: Document, profile: dict[str, Any], data: dict[str, Any]) -> None:
    frame = profile.get("frame") or {}
    fields = frame.get("stamp_fields") if isinstance(frame, dict) else None
    if not isinstance(fields, dict):
        fields = {
            "designation": "{{ system.topic_code }}",
            "title": "{{ system.name }}",
            "developer": "{{ parties.developer }}",
        }
    designation = _render_stamp_field(str(fields.get("designation", "")), data)
    title = _render_stamp_field(str(fields.get("title", "")), data)
    developer = _render_stamp_field(str(fields.get("developer", "")), data)

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear default empty paragraph content lightly
    table = footer.add_table(rows=2, cols=3, width=Cm(16))
    table.style = "Table Grid"
    cells = [
        (0, 0, "Обозначение", designation),
        (0, 1, "Наименование", title),
        (0, 2, "Разраб.", developer),
        (1, 0, "Лист", ""),
        (1, 1, "Листов", ""),
        (1, 2, "", ""),
    ]
    for r, c, label, value in cells:
        cell = table.cell(r, c)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(f"{label}: {value}".strip(": ") if value else label)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_frame(
    doc: Document,
    profile: dict[str, Any],
    data: dict[str, Any] | None = None,
) -> list[str]:
    """Mutate doc according to profile['frame']. Returns warning messages."""
    warnings: list[str] = []
    frame = profile.get("frame") or {}
    if not isinstance(frame, dict):
        return warnings
    preset = str(frame.get("preset") or "none").strip() or "none"
    if preset not in KNOWN_PRESETS:
        warnings.append(f"Неизвестный frame.preset={preset!r}, использован none")
        return warnings
    if preset == "none":
        return warnings
    if preset == "eskd_2_2a":
        warnings.append("Пресет eskd_2_2a пока использует stamp_compact")
        preset = "stamp_compact"

    section = doc.sections[0]
    _set_page_borders(section)
    if preset == "stamp_compact":
        _add_compact_stamp(doc, profile, data or {})
    return warnings
