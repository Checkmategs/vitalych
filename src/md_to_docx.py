"""Convert simple Markdown to DOCX using style-profile.yaml."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.style_profile import alignment, cm, hex_to_rgb, load_style_profile, pt

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|_[^_]+_|`[^`]+`)")


def _set_run_font(
    run,
    *,
    font_family: str,
    size_pt: float,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = font_family
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = hex_to_rgb(color)
    r = run._element.get_or_add_rPr()
    rFonts = r.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_family)
    rFonts.set(qn("w:hAnsi"), font_family)
    rFonts.set(qn("w:eastAsia"), font_family)


def _add_inline_runs(
    paragraph,
    text: str,
    *,
    font_family: str,
    size_pt: float,
    color: str | None = None,
    force_bold: bool | None = None,
) -> None:
    if not text:
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            _set_run_font(
                run,
                font_family=font_family,
                size_pt=size_pt,
                bold=bool(force_bold),
                color=color,
            )
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(
                run,
                font_family=font_family,
                size_pt=size_pt,
                bold=True,
                color=color,
            )
        elif token.startswith("_") and token.endswith("_"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(
                run,
                font_family=font_family,
                size_pt=size_pt,
                italic=True,
                bold=bool(force_bold),
                color=color,
            )
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(
                run,
                font_family=font_family,
                size_pt=size_pt,
                bold=bool(force_bold),
                color=color,
            )
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(
            run,
            font_family=font_family,
            size_pt=size_pt,
            bold=bool(force_bold),
            color=color,
        )


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c or "-") for c in cells)


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _apply_page(doc: Document, profile: dict) -> None:
    page = profile.get("page") or {}
    section = doc.sections[0]
    if "width_cm" in page:
        section.page_width = cm(page["width_cm"])
    if "height_cm" in page:
        section.page_height = cm(page["height_cm"])
    margins = page.get("margins_cm") or {}
    if "top" in margins:
        section.top_margin = cm(margins["top"])
    if "right" in margins:
        section.right_margin = cm(margins["right"])
    if "bottom" in margins:
        section.bottom_margin = cm(margins["bottom"])
    if "left" in margins:
        section.left_margin = cm(margins["left"])
    if "header_distance_cm" in page:
        section.header_distance = cm(page["header_distance_cm"])
    if "footer_distance_cm" in page:
        section.footer_distance = cm(page["footer_distance_cm"])


def _apply_normal_style(doc: Document, profile: dict) -> None:
    body = profile.get("body") or {}
    style = doc.styles["Normal"]
    font_family = body.get("font_family", "Times New Roman")
    size = float(body.get("font_size_pt", 12))
    style.font.name = font_family
    style.font.size = Pt(size)
    if body.get("color"):
        style.font.color.rgb = hex_to_rgb(body["color"])
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_family)
    rFonts.set(qn("w:hAnsi"), font_family)
    rFonts.set(qn("w:eastAsia"), font_family)

    pf = style.paragraph_format
    align = alignment(body.get("alignment"))
    if align is not None:
        pf.alignment = align
    if "left_indent_cm" in body:
        pf.left_indent = cm(body["left_indent_cm"])
    if "right_indent_cm" in body:
        pf.right_indent = cm(body["right_indent_cm"])
    if "first_line_indent_cm" in body:
        pf.first_line_indent = cm(body["first_line_indent_cm"])
    if "space_before_pt" in body:
        pf.space_before = pt(body["space_before_pt"])
    if "space_after_pt" in body:
        pf.space_after = pt(body["space_after_pt"])
    if "line_spacing" in body:
        pf.line_spacing = float(body["line_spacing"])
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def _apply_paragraph_format(paragraph, fmt: dict, *, first_line: bool = True) -> None:
    pf = paragraph.paragraph_format
    align = alignment(fmt.get("alignment"))
    if align is not None:
        pf.alignment = align
    if "left_indent_cm" in fmt:
        pf.left_indent = cm(fmt["left_indent_cm"])
    if "right_indent_cm" in fmt:
        pf.right_indent = cm(fmt["right_indent_cm"])
    if first_line and "first_line_indent_cm" in fmt:
        pf.first_line_indent = cm(fmt["first_line_indent_cm"])
    if "space_before_pt" in fmt:
        pf.space_before = pt(fmt["space_before_pt"])
    if "space_after_pt" in fmt:
        pf.space_after = pt(fmt["space_after_pt"])
    if "line_spacing" in fmt:
        pf.line_spacing = float(fmt["line_spacing"])
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if fmt.get("keep_with_next"):
        pf.keep_with_next = True
    if fmt.get("page_break_before"):
        pf.page_break_before = True


def _add_page_number_footer(doc: Document, profile: dict) -> None:
    footer_cfg = (profile.get("footer") or {}).get("page_number") or {}
    if not footer_cfg.get("enabled", False):
        return
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for child in list(p._p):
        if child.tag.endswith("}r"):
            p._p.remove(child)

    align = alignment(footer_cfg.get("alignment", "right"))
    if align is not None:
        p.alignment = align

    font_family = footer_cfg.get("font_family", "Times New Roman")
    size_pt = float(footer_cfg.get("font_size_pt", 12))

    run = p.add_run()
    _set_run_font(run, font_family=font_family, size_pt=size_pt)

    # PAGE field: begin / instrText / separate / end
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = p.add_run()
    _set_run_font(run2, font_family=font_family, size_pt=size_pt)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = p.add_run()
    _set_run_font(run3, font_family=font_family, size_pt=size_pt)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = p.add_run("1")
    _set_run_font(run4, font_family=font_family, size_pt=size_pt)

    run5 = p.add_run()
    _set_run_font(run5, font_family=font_family, size_pt=size_pt)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def _set_cell_borders(table, color: str, width_pt: float) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    sz = str(int(round(width_pt * 8)))  # eighths of a point
    color_hex = color.lstrip("#").upper()
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)


class _DocxBuilder:
    def __init__(self, profile: dict):
        self.profile = profile
        self.body = profile.get("body") or {}
        self.lists = profile.get("lists") or {}
        self.tables = profile.get("tables") or {}
        levels = (profile.get("headings") or {}).get("levels") or {}
        # YAML may load keys as int
        self.heading_levels = {int(k): v for k, v in levels.items()}
        self.doc = Document()
        self._seen_heading1 = False
        _apply_page(self.doc, profile)
        _apply_normal_style(self.doc, profile)
        _add_page_number_footer(self.doc, profile)

    def add_heading(self, text: str, level: int) -> None:
        level = min(max(level, 1), 5)
        cfg = dict(self.heading_levels.get(level) or {})
        # First H1 should not force page break
        if level == 1:
            if not self._seen_heading1:
                cfg["page_break_before"] = False
                self._seen_heading1 = True
        p = self.doc.add_paragraph()
        style_name = f"Heading {min(level, 3)}"
        try:
            p.style = self.doc.styles[style_name]
        except KeyError:
            pass
        _apply_paragraph_format(p, cfg, first_line=False)
        p.paragraph_format.first_line_indent = cm(0)

        font_family = cfg.get("font_family", self.body.get("font_family", "Times New Roman"))
        size_pt = float(cfg.get("font_size_pt", 12))
        bold = bool(cfg.get("bold", True))
        run = p.add_run(text)
        _set_run_font(run, font_family=font_family, size_pt=size_pt, bold=bold)

    def add_paragraph(self, text: str) -> None:
        p = self.doc.add_paragraph()
        _apply_paragraph_format(p, self.body, first_line=True)
        font_family = self.body.get("font_family", "Times New Roman")
        size_pt = float(self.body.get("font_size_pt", 12))
        color = self.body.get("color")
        _add_inline_runs(
            p,
            text,
            font_family=font_family,
            size_pt=size_pt,
            color=color,
        )

    def add_list_item(self, text: str) -> None:
        p = self.doc.add_paragraph()
        cfg = self.lists
        _apply_paragraph_format(p, cfg, first_line=False)
        left = float(cfg.get("left_indent_cm", 2.521))
        hanging = float(cfg.get("hanging_indent_cm", 0.635))
        p.paragraph_format.left_indent = cm(left)
        p.paragraph_format.first_line_indent = cm(-hanging)

        marker = cfg.get("bullet_marker", "-")
        font_family = cfg.get("font_family", "Times New Roman")
        size_pt = float(cfg.get("font_size_pt", 12))
        _add_inline_runs(
            p,
            f"{marker} {text}",
            font_family=font_family,
            size_pt=size_pt,
            color=self.body.get("color"),
        )

    def add_table(self, rows: list[list[str]]) -> None:
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        cfg = self.tables
        _set_cell_borders(
            table,
            cfg.get("border_color", "#000000"),
            float(cfg.get("border_width_pt", 0.5)),
        )
        widths = cfg.get("preferred_column_widths_cm") or []
        if len(widths) == cols:
            for row in table.rows:
                for j, cell in enumerate(row.cells):
                    cell.width = cm(widths[j])

        font_family = cfg.get("font_family", "Times New Roman")
        size_pt = float(cfg.get("font_size_pt", 12))
        header_bold = bool(cfg.get("header_bold", True))
        header_align = alignment(cfg.get("header_alignment", "center"))
        cell_align = alignment(cfg.get("cell_alignment", "left"))

        for i, row in enumerate(rows):
            for j in range(cols):
                cell = table.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                val = row[j] if j < len(row) else ""
                if i == 0:
                    if header_align is not None:
                        p.alignment = header_align
                    _add_inline_runs(
                        p,
                        val,
                        font_family=font_family,
                        size_pt=size_pt,
                        force_bold=True if header_bold else None,
                    )
                    if header_bold:
                        for run in p.runs:
                            run.bold = True
                else:
                    if cell_align is not None:
                        p.alignment = cell_align
                    _add_inline_runs(
                        p,
                        val,
                        font_family=font_family,
                        size_pt=size_pt,
                    )

    def add_separator(self) -> None:
        p = self.doc.add_paragraph()
        _apply_paragraph_format(p, self.body, first_line=False)
        p.paragraph_format.first_line_indent = cm(0)
        run = p.add_run("─" * 40)
        _set_run_font(
            run,
            font_family=self.body.get("font_family", "Times New Roman"),
            size_pt=10,
            color="#999999",
        )

    def save(self, out_path: Path) -> Path:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out_path))
        return out_path


def markdown_to_docx(
    md_text: str,
    out_path: Path,
    profile: dict | None = None,
    style_profile_path: Path | None = None,
) -> Path:
    if profile is None:
        if style_profile_path is None:
            raise ValueError("profile or style_profile_path is required for DOCX")
        profile = load_style_profile(style_profile_path)

    builder = _DocxBuilder(profile)
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            builder.add_separator()
            i += 1
            continue

        if stripped.startswith("### "):
            builder.add_heading(stripped[4:].strip(), 3)
            i += 1
            continue
        if stripped.startswith("## "):
            builder.add_heading(stripped[3:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("# "):
            builder.add_heading(stripped[2:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            rows = [_split_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not _is_table_sep(lines[i]):
                    rows.append(_split_row(lines[i]))
                i += 1
            builder.add_table(rows)
            continue

        if stripped.startswith("- "):
            builder.add_list_item(stripped[2:].strip())
            i += 1
            continue

        para_parts = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("- ")
                or nxt == "---"
            ):
                break
            para_parts.append(nxt)
            i += 1
        builder.add_paragraph(" ".join(para_parts))

    return builder.save(out_path)
